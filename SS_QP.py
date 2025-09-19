# Generate 3-set Class 6 Social Science question papers with images and export to DOCX

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

import matplotlib.pyplot as plt
import numpy as np
from PIL import Image, ImageDraw, ImageFont
import os
from datetime import date

out_dir = "."
os.makedirs(out_dir, exist_ok=True)

# ---------- Utility: save figure helper ----------
def save_fig(path):
    plt.tight_layout()
    plt.savefig(path, dpi=300, bbox_inches="tight", facecolor='white', edgecolor='none')
    plt.close()

# ---------- Image 1: City map schematic (parametrised) ----------
def generate_city_map(path, seed=0):
    np.random.seed(seed)
    fig, ax = plt.subplots(figsize=(8,6))
    ax.set_aspect('equal')
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis('off')
    fig.patch.set_facecolor('white')
    
    # Roads (grid-like)
    for x in [1.5, 4.5, 7.5]:
        ax.plot([x, x], [0.5, 6.5], linewidth=10, color='#2C3E50')
    for y in [1.5, 3.5, 5.5]:
        ax.plot([0.5, 9.5], [y, y], linewidth=10, color='#2C3E50')
    
    # Landmark positions (avoid overlap by choosing from grid intersections)
    intersections = [(x,y) for x in [1.5,4.5,7.5] for y in [1.5,3.5,5.5]]
    np.random.shuffle(intersections)
    labels = [("Railway Station","RS"), ("Hospital","H"), ("School","SCH"), ("Bank","B"), 
              ("Museum","M"), ("Nagar Panchayat","NP"), ("Public Garden","PG"), ("Apartment","APT")]
    chosen = intersections[:len(labels)]
    
    for (name, short), (x,y) in zip(labels, chosen):
        ax.scatter([x],[y], s=500, marker='s', color='#E74C3C', edgecolor='black', linewidth=2)
        ax.text(x, y, short, ha='center', va='center', fontsize=12, color='white', fontweight='bold')
    
    # Water body
    ax.add_patch(plt.Circle((8.6, 0.9), 0.6, color='#3498DB', alpha=0.8))
    ax.text(8.6, 0.9, "Lake", ha='center', va='center', fontsize=10, color='white', fontweight='bold')
    
    # Compass
    ax.arrow(9.2, 6.0, 0, 0.8, head_width=0.15, head_length=0.25, length_includes_head=True, color='#E67E22', linewidth=3)
    ax.text(9.2, 6.95, "N", ha='center', va='center', fontsize=12, fontweight='bold')
    ax.text(9.2, 5.6, "Compass", ha='center', va='center', fontsize=9, fontweight='bold')
    
    # Legend
    ax.text(0.6, 6.7, "Legend: RS=Railway Station, H=Hospital, SCH=School, B=Bank,\nM=Museum, NP=Nagar Panchayat, PG=Public Garden, APT=Apartment",
            fontsize=9, va='top', bbox=dict(boxstyle="round,pad=0.3", facecolor='lightgray', alpha=0.8))
    save_fig(path)

# ---------- Image 2: Time zone & IDL schematic ----------
def generate_timezones(path):
    fig, ax = plt.subplots(figsize=(8,5))
    ax.axis('off')
    fig.patch.set_facecolor('white')
    # Draw simplified world rectangle
    ax.add_patch(plt.Rectangle((0,0), 360, 180, fill=False, linewidth=2, edgecolor='#2C3E50'))
    # Time zones every 15 degrees
    for lon in range(0, 360, 15):
        ax.plot([lon, lon], [0,180], linewidth=1, color='#7F8C8D', alpha=0.7)
    # Label Greenwich (0°)
    ax.plot([0,0],[0,180], linewidth=3, color='#E74C3C')
    ax.text(2, 172, "0° (Greenwich)", fontsize=10, rotation=90, va='top', fontweight='bold')
    # Label IST (82.5°E)
    ist = 82.5
    ax.plot([ist, ist],[0,180], linewidth=3, color='#27AE60')
    ax.text(ist+2, 172, "82.5°E (IST)", fontsize=10, rotation=90, va='top', fontweight='bold')
    # IDL (approx 180°, wavy)
    x = np.linspace(175, 185, 200)
    y = 90 + 50*np.sin(np.linspace(0, 3*np.pi, 200))
    ax.plot(x, y, linewidth=3, color='#8E44AD')
    ax.text(182, 172, "International Date Line (~180°)", fontsize=10, rotation=90, va='top', fontweight='bold')
    # Direction labels
    ax.text(180, -5, "West (− hours)  ←  Time zones  →  East (+ hours)", ha='center', fontsize=11, fontweight='bold')
    # Gain/lose a day
    ax.annotate("Crossing east → west: +1 day\nCrossing west → east: −1 day",
                xy=(180, 90), xytext=(230, 150),
                arrowprops=dict(arrowstyle="->"))
    # Equator
    ax.plot([0,360],[90,90], linewidth=2, color='#F39C12', linestyle='--')
    ax.text(1, 92, "Equator", fontsize=10, fontweight='bold')
    save_fig(path)

# ---------- Image 3: Archaeological vs Literary sources collage ----------
def generate_sources_collage(path):
    img = Image.new("RGB", (1000, 650), "white")
    d = ImageDraw.Draw(img)
    # Simple icons
    # Coin
    d.ellipse((70,70,220,220), outline="black", width=4)
    d.text((90,230), "Coin", fill="black")
    # Pottery
    d.rectangle((300,120,420,240), outline="black", width=4)
    d.polygon([(360,70),(330,120),(390,120)], outline="black", width=4)
    d.text((320,250), "Pottery", fill="black")
    # Sculpture
    d.rectangle((520,60,640,240), outline="black", width=4)
    d.text((530,250), "Sculpture", fill="black")
    # Palm-leaf manuscript
    d.rectangle((740,80,940,160), outline="black", width=4)
    for x in range(750, 930, 20):
        d.line((x,100,x,140), fill="black")
    d.text((760,170), "Palm-leaf manuscript", fill="black")
    # Inscriptions
    d.rectangle((90,350,320,430), outline="black", width=4)
    for x in range(100, 310, 15):
        d.line((x,370, x+8, 390), fill="black", width=2)
    d.text((110,440), "Inscription", fill="black")
    # Diary/book
    d.rectangle((380,340,520,460), outline="black", width=4)
    d.text((395,470), "Book/Diary", fill="black")
    # Map
    d.rectangle((600,340,900,520), outline="black", width=4)
    d.line((600,430,900,430), fill="black", width=2)
    d.text((690,525), "Map", fill="black")
    # Titles
    d.text((70,15), "Archaeological & Literary Sources (Collage)", fill="black")
    img.save(path, format="PNG")

# ---------- Image 4: Family tree diagram ----------
def generate_family_tree(path):
    img = Image.new("RGB", (1000, 600), "white")
    d = ImageDraw.Draw(img)
    
    # Boxes
    def box(x,y,w,h,text):
        d.rectangle((x,y,x+w,y+h), outline="black", width=3)
        d.text((x+10, y+10), text, fill="black")
    
    # Generation 1
    box(350, 40, 120, 50, "Grandfather")
    box(530, 40, 120, 50, "Grandmother")
    d.line((470,90,590,90), fill="black", width=3)
    d.line((530,90,530,140), fill="black", width=3)
    
    # Generation 2
    box(250, 140, 120, 50, "Uncle")
    box(410, 140, 120, 50, "Father")
    box(570, 140, 120, 50, "Mother")
    box(730, 140, 120, 50, "Aunt")
    # Connectors
    d.line((310,190,310,230), fill="black", width=3)
    d.line((470,190,470,230), fill="black", width=3)
    d.line((630,190,630,230), fill="black", width=3)
    d.line((790,190,790,230), fill="black", width=3)
    
    # Generation 3
    box(220, 230, 120, 50, "Cousin A")
    box(360, 230, 120, 50, "You")
    box(500, 230, 120, 50, "Sibling")
    box(640, 230, 120, 50, "Cousin B")
    # Extra connectors
    d.line((280,280,280,330), fill="black", width=3)
    d.line((420,280,420,330), fill="black", width=3)
    d.line((560,280,560,330), fill="black", width=3)
    d.line((700,280,700,330), fill="black", width=3)
    
    # Roles row (for chores/roles)
    box(200, 330, 160, 50, "Chores: Garden")
    box(380, 330, 160, 50, "Studies: Homework")
    box(560, 330, 160, 50, "Care: Grandparents")
    box(740, 330, 160, 50, "Festivals: Plan")
    img.save(path, format="PNG")

# ---------- Image 5: Timeline template ----------
def generate_timeline(path, start=-8000, end=2025):
    fig, ax = plt.subplots(figsize=(7, 2))
    ax.set_ylim(0,1)
    ax.set_xlim(start, end)
    ax.get_yaxis().set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['top'].set_visible(False)
    ax.spines['bottom'].set_position(('data', 0.5))
    ax.hlines(0.5, start, end, linewidth=2)
    
    ticks = [-8000, -4000, -304, 0, 318, 1858, 2025]
    labels = ["8000 BCE", "4000 BCE", "304 BCE", "0", "318 CE", "1858 CE", "2025 CE"]
    ax.vlines(ticks, 0.45, 0.55, linewidth=2)
    for t, lab in zip(ticks, labels):
        ax.text(t, 0.58, lab, rotation=0, ha='center', va='bottom', fontsize=8)
    ax.text((start+end)/2, 0.85, "Blank Timeline (mark events here)", ha='center', fontsize=10)
    save_fig(path)

# ---------- Generate all images ----------
img_city1 = os.path.join(out_dir, "city_map_set1.png")
img_city2 = os.path.join(out_dir, "city_map_set2.png")
img_city3 = os.path.join(out_dir, "city_map_set3.png")
img_tz1 = os.path.join(out_dir, "time_zones_set1.png")
img_tz2 = os.path.join(out_dir, "time_zones_set2.png")
img_collage = os.path.join(out_dir, "sources_collage.png")
img_family = os.path.join(out_dir, "family_tree.png")
img_timeline = os.path.join(out_dir, "timeline_template.png")

generate_city_map(img_city1, seed=1)
generate_city_map(img_city2, seed=7)
generate_city_map(img_city3, seed=19)
generate_timezones(img_tz1)
generate_timezones(img_tz2)
generate_sources_collage(img_collage)
generate_family_tree(img_family)
generate_timeline(img_timeline)

# ---------- Build the DOCX ----------
doc = Document()

def add_header(doc, set_name):
    title = doc.add_paragraph()
    run = title.add_run(f"Class : 6        HALF YEARLY EXAMINATION (2025-26)\nSubject : SOCIAL SCIENCE")
    run.bold = True
    title_format = title.paragraph_format
    title_format.space_after = Pt(6)
    
    meta = doc.add_paragraph(f"Marks : 80        Time : 3 Hrs.        Set: {set_name}")
    meta_format = meta.paragraph_format
    meta_format.space_after = Pt(12)

def add_section(doc, label, text):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}. {text}")
    run.bold = True

def add_image(doc, path, caption=None, width_inches=5.5):
    doc.add_picture(path, width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph(caption)
        cap_format = cap.paragraph_format
        cap_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_list(doc, items, style="List Number"):
    for it in items:
        doc.add_paragraph(it, style=style)

def add_numbered_list(doc, items):
    for it in items:
        q_num = get_next_question()
        doc.add_paragraph(f"{q_num}. {it}")

def add_image_with_counter(doc, path, caption=None, width_inches=5.5):
    img_num = get_next_image()
    doc.add_picture(path, width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph(f"Picture {chr(64 + img_num)}: {caption}")
        cap_format = cap.paragraph_format
        cap_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

today = date.today().strftime("%d-%m-%Y")

# Global question counter
question_counter = 0
image_counter = 0

def get_next_question():
    global question_counter
    question_counter += 1
    return question_counter

def get_next_image():
    global image_counter
    image_counter += 1
    return image_counter

def reset_counters():
    global question_counter, image_counter
    question_counter = 0
    image_counter = 0

# ----- SET A -----
add_header(doc, "A")

add_section(doc, "I", "Fill in the blanks. [1×8=8]")
add_numbered_list(doc,[
    "The years after the birth of Jesus are counted as _________, and the years before are _________.",
    "The Earth is often called the '__________ planet' because most of its surface is covered with water.",
    "A book or collection of maps is called an _________.",
    "The five major oceans are the Pacific, Atlantic, Indian, Southern and ________ Ocean.",
    "A large unstitched length of cloth worn across India in many styles is called a _________.",
    "In a map, the three important components are distance, direction and _________.",
    "An activity done for money is called an _________ activity.",
    "A family with several generations living together is called a _________ family."
])

add_section(doc, "II", "Mention one difference between the following. [2×6=12]")
add_numbered_list(doc,[
    "Latitude and Longitude",
    "Economic activity and Non-economic activity",
    "Joint family and Nuclear family",
    "Ocean and Sea",
    "Archaeological source and Literary source",
    "Rotation and Revolution"
])

add_section(doc, "III", "Answer as instructed. [2×10=20]")
add_numbered_list(doc,[
    "Define a timeline. Why is there no year zero on the Gregorian calendar?",
    "Arrange these events on the given timeline (next page): Beginning of agriculture (8000 BCE), Beginning of copper metallurgy (4000 BCE), Birth of Aśoka (304 BCE), Birth of Samudragupta (318 CE), Death of Rani Lakshmi Bai (1858 CE).",
    "Calculate how many years before the birth of Christ Aśoka was born.",
    "Name any three staple grains commonly used across India.",
    "Mention any two reasons why oceans are vital for life on Earth.",
    "Give two features of a physical map and two of a political map.",
    "State two benefits of living in a community.",
    "Name the levels of government in India.",
    "Give two examples of standard symbols commonly used in maps.",
    "Classify the following as economic or non-economic: (a) Teaching at a paid tuition class (b) Cooking dinner for the family."
])

add_image_with_counter(doc, img_timeline, caption="Blank timeline for Q. 18. Fill and label appropriately.", width_inches=6.0)

add_section(doc, "IV", "Observe the pictures and answer the questions. [1×15=15]")
add_image_with_counter(doc, img_city1, caption="Schematic city map")
add_numbered_list(doc,[
    "Define a map.",
    "Which important component of a map indicates the cardinal directions?",
    "From the Railway Station (RS), which direction is the Hospital (H)?",
    "Suggest the shortest road route from the School (SCH) to the Apartment (APT).",
    "What does the cross symbol (H) usually denote on maps?",
    "Why do we use standard symbols on maps?",
    "Estimate the longitudinal (left–right) extent of the Museum (M) on this map (qualitative answer).",
    "Name one instrument used by sailors in ancient times to navigate.",
    "Name the modern system widely used for navigation today.",
    "If each large square on the map equals 1 km, approximately how far is the Bank (B) from the Public Garden (PG)?",
    "State two differences between a plan and a map.",
    "Why is a scale important in maps?",
    "If you walk from the Nagar Panchayat (NP) due east to the next road, which landmark do you reach first?",
    "Identify one possible hazard near the 'Lake' and suggest a safety sign for it.",
    "Add a legend entry for the Apartment using an appropriate symbol."
])

add_image_with_counter(doc, img_tz1, caption="Time zones and the International Date Line")

add_numbered_list(doc,[
    "Define a time zone.",
    "What is the significance of the International Date Line?",
    "If you travel from the USA to Australia across the IDL, will you gain or lose a day? Explain briefly.",
    "Why are time zone boundaries not perfectly straight?",
    "At which longitude is Indian Standard Time (IST) calculated?"
])

add_section(doc, "V", "Answer the following in 3–4 sentences. [3×5=15]")
add_numbered_list(doc,[
    "Explain why India is often called 'Bhārat'. Mention one ancient source that uses this name.",
    "How do oceans and continents together influence climate and human life?",
    "Explain representative democracy and name the three organs of government.",
    "How can family roles and responsibilities change over time? Give examples.",
    "Explain the idea of 'unity in diversity' with one cultural example."
])

add_section(doc, "VI", "Answer the following in 6–7 sentences. (Any two) [5×2=10]")
add_numbered_list(doc,[
    "Describe the functions of the Legislature, Executive and Judiciary with one example each.",
    "Differentiate clearly between archaeological and literary sources of history with suitable examples.",
    "Explain how latitude and longitude help us locate places on the Earth; include an example using coordinates."
])

doc.add_page_break()

# Reset counters for Set B
reset_counters()

# ----- SET B -----
add_header(doc, "B")

add_section(doc, "I", "Fill in the blanks. [1×8=8]")
add_numbered_list(doc,[
    "The imaginary line dividing the Earth into Northern and Southern Hemispheres is the _________.",
    "The conventional year for the birth of Jesus is taken as the start of the _________ Era.",
    "Standard time in India is calculated at _________°E longitude.",
    "A large body of continuous land is called a _________.",
    "Work done out of love or duty without payment is a _________ activity.",
    "A family consisting of parents and their children is called a _________ family.",
    "Objects like tools, pots and coins dug out from the ground are studied by _________.",
    "A book of maps is called an _________."
])

add_section(doc, "II", "Mention one difference between the following. [2×6=12]")
add_numbered_list(doc,[
    "Primary source and Secondary source of history",
    "Physical map and Political map",
    "Galaxy and Universe",
    "Standard time and Local time",
    "Staple grains and Spices",
    "Revolution and Rotation"
])

add_section(doc, "III", "Answer as instructed. [2×10=20]")
add_numbered_list(doc,[
    "Explain, with an example, how to calculate the number of years between a BCE date and a CE date.",
    "Name the five oceans of the world.",
    "List any four standard map symbols you have studied.",
    "State any two benefits of living in a community.",
    "Give two reasons why freshwater is precious despite the abundance of seawater.",
    "Name any three regions or kingdoms listed in the Mahābhārata that cover different parts of the Subcontinent.",
    "Classify each of the following as economic/non-economic: (a) Driving a taxi (b) Caring for a sick grandparent (c) Selling vegetables in a market (d) Volunteering in a beach clean-up.",
    "Define 'Era'. Give the names of two calendar eras commonly used in India.",
    "Name two features common to traditional Indian saris across regions.",
    "State two ways in which oceans support biodiversity."
])

add_section(doc, "IV", "Observe the picture and answer the questions. [1×15=15]")
add_image_with_counter(doc, img_collage, caption="Collage of sources")
add_numbered_list(doc,[
    "Identify any three archaeological sources visible in the collage.",
    "Identify any two literary sources visible in the collage.",
    "Define 'archaeology'.",
    "Why are inscriptions useful to historians?",
    "Give one example each of what coins and pottery can tell us about the past."
])

add_image_with_counter(doc, img_city2, caption="Schematic city map")
add_numbered_list(doc,[
    "From the School (SCH), which direction is the Bank (B)?",
    "Suggest the quickest route from the Railway Station (RS) to the Museum (M).",
    "Add an appropriate symbol in the legend for 'Temple' and mark a possible location on the map.",
    "Why is a compass important on maps?",
    "What is the purpose of a map scale? Give one example."
])

add_section(doc, "V", "Answer the following in 3–4 sentences. [3×5=15]")
add_numbered_list(doc,[
    "Explain how oceans are conventionally divided and why those divisions are not natural boundaries.",
    "Describe two ways families show cooperation and interdependence.",
    "What do you understand by 'subcontinent'? Name two ancient terms used for India."
])

add_section(doc, "VI", "Answer the following in 6–7 sentences. (Any two) [5×2=10]")
add_numbered_list(doc,[
    "Discuss with examples how the sari illustrates India's 'unity in diversity'.",
    "Explain the relation between longitude and time with a clear example of time difference.",
    "Describe the role of standard symbols in making maps readable and comparable."
])

doc.add_page_break()

# Reset counters for Set C
reset_counters()

# ----- SET C -----
add_header(doc, "C")

add_section(doc, "I", "Fill in the blanks. [1×8=8]")
add_numbered_list(doc,[
    "The smallest of the five oceans is the _________ Ocean.",
    "The line of 0° longitude is also called the _________ Meridian.",
    "The Indian name 'Bhāratavarṣa' appears in ancient texts such as the _________.",
    "The discipline that studies human societies and cultures from the oldest times to the present is called _________.",
    "A large expanse of water smaller than an ocean but partly enclosed by land may be called a _________.",
    "The three main organs of the Indian government are the Legislature, _________ and _________.",
    "A _________ is a convenient tool to mark events in chronological order.",
    "The _________ Date Line roughly follows the 180° longitude with deviations."
])

add_section(doc, "II", "Mention one difference between the following. [2×6=12]")
add_numbered_list(doc,[
    "Continent and Island",
    "Archaeologist and Anthropologist",
    "Joint family and Nuclear family",
    "Map and Globe",
    "Economic activity and Service (sevā)",
    "Local Government and Central Government"
])

add_section(doc, "III", "Answer as instructed. [2×10=20]")
add_numbered_list(doc,[
    "Define 'standard time'. Why does a large country like India use one standard time?",
    "Name three calendars (other than the Gregorian) that are commonly used for festivals in India.",
    "List any four staple grains or pulses widely used in India.",
    "Write two reasons why symbols and legends are essential parts of a map.",
    "Name the five oceans and identify the largest among them.",
    "Explain with numbers how 2 BCE to 2 CE spans only three years on a timeline.",
    "Mention any two ways oceans affect climate.",
    "State two examples to show changing roles and responsibilities within a family over time.",
    "Give two differences between rotation and revolution of the Earth.",
    "What is meant by 'subcontinent'? Name one physical feature that helps define India's natural boundaries."
])

add_section(doc, "IV", "Observe the pictures and answer the questions. [1×15=15]")
add_image_with_counter(doc, img_family, caption="Family Tree (schematic)")
add_numbered_list(doc,[
    "Name the two main types of families discussed in your textbook and identify which type best fits the picture.",
    "Suggest two responsibilities you can take up to support your family.",
    "How can elderly family members contribute to children's learning? Give one example.",
    "Explain how cooperation within a family helps during festivals or emergencies.",
    "Give one example of sharing resources in a joint family."
])

add_image_with_counter(doc, img_tz2, caption="Time zones and IDL")
add_numbered_list(doc,[
    "Define 'longitude'.",
    "What is the longitude used for IST?",
    "If it is 12:00 noon at 0°, what is the approximate time at 82.5°E? (Show the step clearly.)",
    "Why does crossing the IDL change the date?",
    "State one reason why the IDL is not a straight line."
])

add_image_with_counter(doc, img_city3, caption="Schematic city map")
add_numbered_list(doc,[
    "From the Bank (B), in which direction is the Apartment (APT)?",
    "Add a suitable symbol for 'Bus stop' and place it near a road junction. Explain your choice.",
    "Estimate the distance from the Museum (M) to the Hospital (H) if each big square = 1 km.",
])

add_section(doc, "V", "Answer the following in 3–4 sentences. [3×5=15]")
add_numbered_list(doc,[
    "Explain how historians use inscriptions and coins to reconstruct the past.",
    "What does the phrase 'unity in diversity' mean? Support your answer with one cultural example.",
    "Describe why freshwater is scarce and list one way to conserve it at home or school."
])

add_section(doc, "VI", "Answer the following in 6–7 sentences. (Any two) [5×2=10]")
add_numbered_list(doc,[
    "Describe India's ancient names and what they reveal about the idea of India.",
    "Explain how latitudes and longitudes together help to locate a place on the globe with an example of coordinates.",
    "Discuss the social and economic value of both paid work and unpaid service (sevā) in a community."
])

# Footer note
doc.add_paragraph("\nNOTE: Images are schematic and intended for educational use in exam questions.")

out_path = os.path.join(out_dir, "Class6_SST_HalfYearly_3Sets.docx")
doc.save(out_path)

out_path
